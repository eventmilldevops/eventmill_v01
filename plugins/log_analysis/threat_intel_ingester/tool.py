"""
Threat Intel Ingester — Event Mill Plugin

Ingests threat intelligence reports (PDF, HTML, STIX, CSV/JSON IOC lists)
and extracts structured IOC data with MITRE ATT&CK mapping.

This plugin depends on LLM capabilities for contextual IOC extraction
and MITRE technique inference. Regex-based extraction provides a baseline;
LLM analysis provides confidence scoring and priority assessment.
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from framework.logging.structured import log_llm_interaction
from framework.plugins.protocol import ToolResult, ValidationResult, QueryHints
from framework.reference_data.mitre_attack import get_mitre_db as _get_mitre_db

logger = logging.getLogger("eventmill.plugin.threat_intel_ingester")

# ---------------------------------------------------------------------------
# IOC Regex Patterns
# ---------------------------------------------------------------------------

IOC_PATTERNS = {
    "ip": re.compile(
        r"\b(?:(?:25[0-5]|2[0-4]\d|1?\d\d?)(?:\.|\[\.\]))"
        r"{3}(?:25[0-5]|2[0-4]\d|1?\d\d?)\b"
    ),
    "domain": re.compile(
        r"\b(?:[a-zA-Z0-9](?:[a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?"
        r"(?:\.|\[\.\]))+(?:com|net|org|io|info|biz|xyz|top|"
        r"ru|cn|de|uk|fr|jp|br|au|ca|nl|it|es|ch|se|no|fi|"
        r"dk|be|at|pl|cz|sk|hu|ro|bg|hr|si|lt|lv|ee|ie|pt|"
        r"gr|cy|lu|mt|li|is)\b",
        re.IGNORECASE,
    ),
    "hash_md5": re.compile(r"\b[a-fA-F0-9]{32}\b"),
    "hash_sha1": re.compile(r"\b[a-fA-F0-9]{40}\b"),
    "hash_sha256": re.compile(r"\b[a-fA-F0-9]{64}\b"),
    "url": re.compile(
        r"(?:https?|hxxps?|ftp)(?:://|(?:\[:\]//))[\w\-._~:/?#\[\]@!$&'()*+,;=%]+",
        re.IGNORECASE,
    ),
    "email": re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Z|a-z]{2,}\b"),
    "cve": re.compile(r"\bCVE-\d{4}-\d{4,7}\b", re.IGNORECASE),
    "mitre_technique": re.compile(r"\bT\d{4}(?:\.\d{3})?\b"),
}

# Defanging reversal patterns
DEFANG_REPLACEMENTS = [
    (re.compile(r"\[\.\]"), "."),
    (re.compile(r"hxxp", re.IGNORECASE), "http"),
    (re.compile(r"\[:\]"), ":"),
    (re.compile(r"\[at\]", re.IGNORECASE), "@"),
]


def refang(value: str) -> str:
    """Reverse common defanging patterns."""
    result = value
    for pattern, replacement in DEFANG_REPLACEMENTS:
        result = pattern.sub(replacement, result)
    return result


def was_defanged(original: str, refanged: str) -> bool:
    """Check if the value was defanged in the original text."""
    return original != refanged


# ---------------------------------------------------------------------------
# MITRE ATT&CK Kill-Chain Ordering
# ---------------------------------------------------------------------------

TACTIC_ORDER: dict[str, int] = {
    "Reconnaissance": 1,
    "Resource Development": 2,
    "Initial Access": 3,
    "Execution": 4,
    "Persistence": 5,
    "Privilege Escalation": 6,
    "Defense Evasion": 7,
    "Credential Access": 8,
    "Discovery": 9,
    "Lateral Movement": 10,
    "Collection": 11,
    "Command and Control": 12,
    "Exfiltration": 13,
    "Impact": 14,
}

# Tactics that should only appear at the entry point (first step) of a path.
# If a later step is assigned one of these and the technique has alternatives,
# the reconciler will reassign using the kill-chain ordering.
ENTRY_ONLY_TACTICS: set[str] = {"Reconnaissance", "Resource Development", "Initial Access"}


# ---------------------------------------------------------------------------
# Text Extraction Helpers
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_path: str, max_pages: int = 50) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("pdfplumber is required for PDF processing")

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            if i >= max_pages:
                break
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_html(file_path: str) -> str:
    """Extract text from an HTML file using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError("beautifulsoup4 is required for HTML processing")

    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    # Remove script and style elements
    for element in soup(["script", "style", "nav", "footer", "header"]):
        element.decompose()

    return soup.get_text(separator="\n", strip=True)


def extract_text_from_text(file_path: str) -> str:
    """Read a plain text file (including Markdown)."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is required for Word document processing. "
            "Install with: pip install python-docx"
        )
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


TEXT_EXTRACTORS = {
    "pdf_report": extract_text_from_pdf,
    "html_report": extract_text_from_html,
    "text": extract_text_from_text,
    "docx_report": extract_text_from_docx,
}


# ---------------------------------------------------------------------------
# Regex-Based IOC Extraction
# ---------------------------------------------------------------------------


@dataclass
class RawIOC:
    """An IOC extracted by regex before LLM refinement."""
    ioc_type: str
    value: str
    raw_value: str
    defanged: bool
    context: str = ""
    confidence: str = "low"
    priority: str = "medium"
    related_mitre: list[str] = field(default_factory=list)


def extract_iocs_regex(
    text: str,
    ioc_types: list[str],
) -> list[RawIOC]:
    """Extract IOCs from text using regex patterns.

    Returns deduplicated IOCs with surrounding context.
    """
    seen: set[tuple[str, str]] = set()
    results: list[RawIOC] = []

    for ioc_type in ioc_types:
        pattern = IOC_PATTERNS.get(ioc_type)
        if not pattern:
            continue

        for match in pattern.finditer(text):
            raw_value = match.group(0)
            value = refang(raw_value)
            defanged = was_defanged(raw_value, value)

            key = (ioc_type, value.lower())
            if key in seen:
                continue
            seen.add(key)

            # Extract surrounding context (up to 150 chars each side)
            start = max(0, match.start() - 150)
            end = min(len(text), match.end() + 150)
            context = text[start:end].replace("\n", " ").strip()

            results.append(
                RawIOC(
                    ioc_type=ioc_type,
                    value=value,
                    raw_value=raw_value,
                    defanged=defanged,
                    context=context[:300],
                    confidence="low",
                )
            )

    return results


# ---------------------------------------------------------------------------
# Chunked LLM processing helpers
# ---------------------------------------------------------------------------

_MAX_IOC_PER_CHUNK: int = 50        # IOC candidates per LLM call
_MAX_TEXT_CHARS_PER_CHUNK: int = 6_000  # Report text chars per LLM call


def _chunk_text(text: str, max_chars: int = _MAX_TEXT_CHARS_PER_CHUNK) -> list[str]:
    """Split text into paragraph-bounded chunks, each under max_chars."""
    if len(text) <= max_chars:
        return [text]
    chunks: list[str] = []
    paragraphs = text.split("\n\n")
    current: list[str] = []
    current_len = 0
    for para in paragraphs:
        para_len = len(para) + 2  # account for the "\n\n" separator
        if current_len + para_len > max_chars and current:
            chunks.append("\n\n".join(current))
            current = [para]
            current_len = para_len
        else:
            current.append(para)
            current_len += para_len
    if current:
        chunks.append("\n\n".join(current))
    return chunks or [text[:max_chars]]


def _merge_llm_chunk_results(chunk_results: list[dict]) -> dict:
    """Merge LLM JSON results from multiple document chunks.

    IOCs are deduplicated by (type, value); MITRE techniques by (technique_id, tactic).
    report_metadata comes from the first successful chunk.
    attack_graph paths are unioned and deduplicated by path_id.
    """
    merged_iocs: list[dict] = []
    merged_mitre: list[dict] = []
    seen_ioc_keys: set[str] = set()
    seen_mitre_keys: set[tuple[str, str]] = set()
    report_meta: dict = {}
    ag_paths: list[dict] = []
    ag_convergence: set[str] = set()
    ag_branches: set[str] = set()

    for result in chunk_results:
        for ioc in result.get("refined_iocs", []):
            key = f"{ioc.get('ioc_type')}:{ioc.get('value', '').lower()}"
            if key not in seen_ioc_keys:
                seen_ioc_keys.add(key)
                merged_iocs.append(ioc)

        for m in result.get("additional_mitre_techniques", []):
            tid = m.get("technique_id", "")
            tactic = m.get("tactic", "")
            mkey = (tid, tactic)
            if tid and mkey not in seen_mitre_keys:
                seen_mitre_keys.add(mkey)
                merged_mitre.append(m)

        if not report_meta:
            report_meta = result.get("report_metadata") or {}

        ag = result.get("attack_graph") or {}
        for path in ag.get("paths", []):
            pid = path.get("path_id")
            if not any(p.get("path_id") == pid for p in ag_paths):
                ag_paths.append(path)
        ag_convergence.update(ag.get("convergence_points", []))
        ag_branches.update(ag.get("branch_points", []))

    return {
        "refined_iocs": merged_iocs,
        "additional_mitre_techniques": merged_mitre,
        "report_metadata": report_meta,
        "attack_graph": {
            "paths": ag_paths,
            "convergence_points": list(ag_convergence),
            "branch_points": list(ag_branches),
        } if ag_paths else {},
    }


def _parse_llm_json(response_text: str) -> dict | None:
    """Strip markdown code fences and parse JSON from LLM response.

    Logs the exact parse error on failure.  If the JSON appears truncated
    (common when the model hits its output-token limit), attempts
    best-effort repair by closing unmatched brackets.
    """
    _log = logging.getLogger("eventmill.plugin.threat_intel_ingester")
    raw_resp = response_text.strip()

    had_fences = raw_resp.startswith("```")
    if had_fences:
        raw_resp = re.sub(r"^```(?:json)?\s*\n?", "", raw_resp)
        raw_resp = re.sub(r"\n?```\s*$", "", raw_resp)

    text = raw_resp.strip()
    resp_len = len(text)

    # --- Fast path: direct parse ---
    try:
        return json.loads(text)
    except json.JSONDecodeError as exc:
        _log.warning(
            "JSON parse error at char %d (line %d col %d): %s "
            "| response_length=%d, had_fences=%s, "
            "last_100_chars=%r",
            exc.pos, exc.lineno, exc.colno, exc.msg,
            resp_len, had_fences,
            text[-100:] if resp_len > 100 else text,
        )
    except ValueError as exc:
        _log.warning(
            "JSON ValueError: %s | response_length=%d",
            exc, resp_len,
        )

    # --- Slow path: repair truncated JSON ---
    repaired = _repair_truncated_json(text)
    if repaired is not None:
        _log.info(
            "Recovered truncated JSON via bracket repair "
            "(original_length=%d, keys=%s)",
            resp_len, list(repaired.keys()),
        )
        return repaired

    _log.warning(
        "JSON repair also failed "
        "| response_length=%d, starts_with_brace=%s",
        resp_len, text[:1] == '{',
    )
    return None


def _repair_truncated_json(text: str) -> dict | None:
    """Best-effort repair of JSON truncated by LLM output-token limits.

    Scans backward from the end of *text* looking for the last ``}`` or
    ``]`` that, when followed by the right number of closing brackets,
    produces a valid JSON object.  Tries up to 30 candidate positions.
    """
    if not text.startswith('{'):
        return None

    attempts = 0
    for i in range(len(text) - 1, max(0, len(text) - 4000), -1):
        if text[i] not in ('}', ']'):
            continue
        if attempts >= 30:
            break
        attempts += 1

        candidate = text[:i + 1]
        need_brackets = candidate.count('[') - candidate.count(']')
        need_braces = candidate.count('{') - candidate.count('}')
        if need_brackets < 0 or need_braces < 0:
            continue

        closer = ']' * need_brackets + '}' * need_braces
        try:
            obj = json.loads(candidate + closer)
            if isinstance(obj, dict):
                return obj
        except (json.JSONDecodeError, ValueError):
            continue

    return None


# ---------------------------------------------------------------------------
# MITRE ATT&CK technique lookup — delegated to framework.reference_data
# ---------------------------------------------------------------------------
# _get_mitre_db is imported above from framework.reference_data.mitre_attack


def _fix_tactic_progression(
    attack_graph: dict,
    mitre_db: dict[str, dict],
) -> tuple[dict, int]:
    """Enforce kill-chain tactic progression within each attack path.

    **Entry-point-only rule**: ``Initial Access``, ``Reconnaissance``, and
    ``Resource Development`` should only appear at the first step of a path.
    If a later step uses one of these and the technique has alternative valid
    tactics in the MITRE database, reassign to the valid tactic with the
    highest kill-chain ordinal (preferring forward progression).

    Mutates the ``attack_graph`` in place and returns it along with the
    count of reassignments made.
    """
    reassign_count = 0

    for path in attack_graph.get("paths", []):
        path_id = path.get("path_id", "unknown")
        steps = path.get("steps", [])
        if len(steps) < 2:
            continue

        for step_idx, step in enumerate(steps):
            if step_idx == 0:
                continue  # first step is allowed to use entry-only tactics

            tid = step.get("technique_id", "")
            tactic = step.get("tactic", "")
            if not tid or not tactic:
                continue

            if tactic not in ENTRY_ONLY_TACTICS:
                continue  # tactic is fine, no fix needed

            # Technique is using an entry-only tactic at a non-first step.
            # Look up its valid tactics from the MITRE database.
            local = mitre_db.get(tid, {})
            valid_tactics: list[str] = local.get("tactics", [])
            if not valid_tactics:
                continue  # no DB data to pick from

            # Filter out entry-only tactics, keep alternatives
            alternatives = [
                t for t in valid_tactics
                if t not in ENTRY_ONLY_TACTICS and t in TACTIC_ORDER
            ]
            if not alternatives:
                continue  # all valid tactics are entry-only; keep as-is

            # Pick the alternative with the highest kill-chain ordinal
            best = max(alternatives, key=lambda t: TACTIC_ORDER[t])

            logger.info(
                "[TACTIC-FIX] Path %r step %d: %s reassigned "
                "%r -> %r (entry-only tactic at non-first step; "
                "valid alternatives: %s)",
                path_id, step_idx, tid, tactic, best, alternatives,
            )
            step["tactic"] = best
            reassign_count += 1

    return attack_graph, reassign_count


def _reconcile_mitre_mappings(
    all_mitre: list[dict],
    attack_graph: dict,
) -> list[dict]:
    """Reconcile and enrich mitre_mappings against local ATT&CK data.

    Identity key is ``(technique_id, tactic)`` — one technique can appear
    multiple times with different tactics when it serves different roles
    across attack paths.

    0. Runs ``_fix_tactic_progression`` on the attack_graph to reassign
       entry-only tactics (Initial Access, etc.) on non-first steps.
    1. Backfills ``(technique_id, tactic)`` pairs from *attack_graph* steps,
       populating ``context_paths`` with the path IDs where each pair appears.
    2. Enriches entries that have empty ``technique_name`` or ``tactic``
       using the local MITRE lookup.
    3. Validates technique IDs, marks non-ATT&CK entries, and flags tactic
       mismatches.

    Returns the (mutated) *all_mitre* list.
    """
    mitre_db = _get_mitre_db()

    # --- Step 0: fix tactic progression in attack_graph ---
    _fix_tactic_progression(attack_graph, mitre_db)

    # --- Index existing entries by (technique_id, tactic) ---
    existing: dict[tuple[str, str], dict] = {}
    for m in all_mitre:
        tid = m.get("technique_id", "")
        tactic = m.get("tactic", "")
        if tid:
            existing[(tid, tactic)] = m

    backfill_count = 0
    enrich_count = 0

    # --- Step 1: backfill from attack_graph with per-step tactic ---
    # Collect all (tid, tactic) -> [path_ids] from graph steps
    graph_keys: dict[tuple[str, str], list[str]] = {}
    all_leads_to: set[str] = set()
    step_tids: set[str] = set()

    for path in attack_graph.get("paths", []):
        path_id = path.get("path_id", "unknown")
        for step in path.get("steps", []):
            tid = step.get("technique_id", "")
            tactic = step.get("tactic", "")
            if not tid:
                continue
            step_tids.add(tid)
            key = (tid, tactic)
            graph_keys.setdefault(key, [])
            if path_id not in graph_keys[key]:
                graph_keys[key].append(path_id)
            all_leads_to.update(t for t in step.get("leads_to", []) if t)

    for key, path_ids in graph_keys.items():
        tid, tactic = key
        if key in existing:
            # Exact (tid, tactic) match — just add context_paths
            entry = existing[key]
            paths_list = entry.setdefault("context_paths", [])
            for pid in path_ids:
                if pid not in paths_list:
                    paths_list.append(pid)
            continue

        # Check for existing entry with same tid but empty tactic
        # (from IOC-derived or LLM additional_mitre with empty tactic)
        empty_key = (tid, "")
        if empty_key in existing:
            entry = existing.pop(empty_key)
            entry["tactic"] = tactic
            paths_list = entry.setdefault("context_paths", [])
            for pid in path_ids:
                if pid not in paths_list:
                    paths_list.append(pid)
            existing[key] = entry
            enrich_count += 1
            logger.info(
                "[RECONCILE] Promoted %s: set tactic=%r from "
                "attack_graph path(s) %s",
                tid, tactic, ", ".join(path_ids),
            )
            continue

        # New (tid, tactic) pair — create entry
        local = mitre_db.get(tid, {})
        new_entry = {
            "technique_id": tid,
            "technique_name": local.get("name", ""),
            "tactic": tactic,
            "context_paths": list(path_ids),
            "confidence": "inferred",
            "report_context": (
                f"Backfilled from attack_graph path(s) "
                f"'{', '.join(path_ids)}'"
            ),
        }
        all_mitre.append(new_entry)
        existing[key] = new_entry
        backfill_count += 1
        logger.info(
            "[RECONCILE] Backfilled technique %s (tactic=%s) "
            "from attack_graph path(s) %s | local_lookup=%s",
            tid, tactic, ", ".join(path_ids),
            "hit" if local else "miss",
        )

    # Handle leads_to targets that never appear as steps in any path
    orphan_tids = all_leads_to - step_tids
    for oid in sorted(orphan_tids):
        # Skip if any entry already exists for this technique_id
        if any(etid == oid for (etid, _) in existing):
            continue
        local = mitre_db.get(oid, {})
        fallback_tactic = local["tactics"][0] if local.get("tactics") else ""
        new_entry = {
            "technique_id": oid,
            "technique_name": local.get("name", ""),
            "tactic": fallback_tactic,
            "context_paths": [],
            "confidence": "inferred",
            "report_context": (
                "Backfilled — referenced as leads_to target "
                "in attack_graph"
            ),
        }
        all_mitre.append(new_entry)
        existing[(oid, fallback_tactic)] = new_entry
        backfill_count += 1
        logger.info(
            "[RECONCILE] Backfilled leads_to orphan %s (tactic=%s) "
            "| local_lookup=%s",
            oid, fallback_tactic or "(unknown)",
            "hit" if local else "miss",
        )

    # --- Step 2: enrich stubs with empty name/tactic ---
    for entry in all_mitre:
        tid = entry.get("technique_id", "")
        if not tid:
            continue

        needs_name = not entry.get("technique_name")
        needs_tactic = not entry.get("tactic")
        if not needs_name and not needs_tactic:
            continue

        local = mitre_db.get(tid, {})
        if not local:
            continue

        changes: list[str] = []
        if needs_name and local.get("name"):
            entry["technique_name"] = local["name"]
            changes.append(f"name={local['name']!r}")
        if needs_tactic and local.get("tactics"):
            if len(local["tactics"]) > 1:
                logger.warning(
                    "[RECONCILE] %s has %d valid tactics %s but no "
                    "graph context to disambiguate — defaulting to %r",
                    tid, len(local["tactics"]), local["tactics"],
                    local["tactics"][0],
                )
            entry["tactic"] = local["tactics"][0]
            changes.append(f"tactic={local['tactics'][0]!r}")

        if changes:
            enrich_count += 1
            logger.info(
                "[RECONCILE] Enriched %s: %s | from local MITRE lookup",
                tid, ", ".join(changes),
            )

    # --- Step 3: validate IDs and mark non-ATT&CK entries ---
    # Only run when we actually have a loaded DB; skip if the lookup
    # file hasn't been built yet (all entries stay unmarked).
    unvalidated_count = 0
    tactic_mismatch_count = 0
    if mitre_db:
        for entry in all_mitre:
            tid = entry.get("technique_id", "")
            if not tid:
                continue
            if tid in mitre_db:
                entry["mitre_validated"] = True
                # Validate tactic against allowed tactics for this technique
                entry_tactic = entry.get("tactic", "")
                allowed = mitre_db[tid].get("tactics", [])
                if entry_tactic and allowed:
                    # Case-insensitive lookup: LLM may write
                    # "Command and Control" vs DB's "Command And Control"
                    allowed_lower = {t.lower(): t for t in allowed}
                    if entry_tactic in allowed:
                        # Exact match — nothing to do
                        pass
                    elif entry_tactic.lower() in allowed_lower:
                        # Case-only mismatch — auto-correct to DB casing
                        canonical = allowed_lower[entry_tactic.lower()]
                        logger.info(
                            "[RECONCILE] Auto-corrected tactic casing for "
                            "%s: %r -> %r",
                            tid, entry_tactic, canonical,
                        )
                        entry["tactic"] = canonical
                    else:
                        # Genuine mismatch — flag in output
                        tactic_mismatch_count += 1
                        entry["tactic_mismatch"] = True
                        logger.warning(
                            "[RECONCILE] Tactic mismatch: %s assigned "
                            "tactic %r but ATT&CK allows %s — "
                            "keeping LLM assignment, flagged in output",
                            tid, entry_tactic, allowed,
                        )
            else:
                entry["mitre_validated"] = False
                unvalidated_count += 1
                # Annotate the name so frontline analysts see it
                # without needing access to tool logs.
                name = entry.get("technique_name", "")
                if name and "(non-ATT&CK ID)" not in name:
                    entry["technique_name"] = f"{name} (non-ATT&CK ID)"
                elif not name:
                    entry["technique_name"] = "(non-ATT&CK ID)"
                logger.warning(
                    "[RECONCILE] Unvalidated technique %s (%s) — "
                    "not found in ATT&CK v18.1 (DB has %d techniques). "
                    "Keeping entry but marking as non-ATT&CK.",
                    tid, entry["technique_name"], len(mitre_db),
                )

    if backfill_count or enrich_count or unvalidated_count or tactic_mismatch_count:
        logger.info(
            "[RECONCILE] Summary: %d backfilled, %d enriched, "
            "%d unvalidated, %d tactic mismatches, "
            "%d total mitre_mappings (local DB has %d techniques)",
            backfill_count, enrich_count, unvalidated_count,
            tactic_mismatch_count,
            len(all_mitre), len(mitre_db),
        )

    return all_mitre


# ---------------------------------------------------------------------------
# LLM Refinement Prompt
# ---------------------------------------------------------------------------

LLM_REFINEMENT_PROMPT = """You are an experienced threat intelligence analyst reviewing extracted IOCs from a security report.

TASK: Complete ALL five sections below.

SECTION 1 — IOC VALIDATION: Review each IOC candidate and its surrounding context.
For each IOC:
1. Assess whether it is a true indicator of compromise (not a benign version number, documentation example, or false positive)
2. Assign a confidence level: "low" (uncertain), "medium" (likely real IOC), "high" (confirmed IOC based on context)
3. Assign an operational priority: "low", "medium", "high" based on the IOC's role in the described attack
4. Identify related MITRE ATT&CK technique IDs if the context suggests a specific technique

SECTION 2 — ADDITIONAL MITRE TECHNIQUES: Identify any MITRE ATT&CK techniques described in the report text that are not already captured as IOC-type extractions. For each technique, note whether it was explicitly mentioned (technique ID appears in text) or inferred from described behavior.

SECTION 3 — REPORT METADATA: Extract the report title, campaign name, attributed threat actor, and attribution confidence.

SECTION 4 — TECHNIQUE TACTIC ASSIGNMENT: For EVERY technique in both refined_iocs.related_mitre AND additional_mitre_techniques, you MUST populate the "tactic" field with the correct MITRE ATT&CK tactic name. Use the official tactic names exactly as written:
Reconnaissance, Resource Development, Initial Access, Execution, Persistence, Privilege Escalation, Defense Evasion, Credential Access, Discovery, Lateral Movement, Collection, Command and Control, Exfiltration, Impact.
If a technique maps to multiple tactics, use the tactic most relevant to how the report describes its use. When the same technique ID appears in multiple attack paths serving different attacker objectives, include it multiple times in `additional_mitre_techniques` — once per distinct role — with the tactic that matches each role. This is expected, not a duplication error. NEVER leave the tactic field empty.

SECTION 5 — ATTACK GRAPH: Analyze how the techniques described in the report relate to each other operationally. Real attacks have multiple paths, branches, and convergence points.

For the attack_graph:
- Identify distinct attack PATHS (e.g., "phishing path" and "supply chain path" may both lead to execution)
- For each path, list the techniques in causal order: which technique ENABLES or LEADS TO the next
- Identify CONVERGENCE POINTS — techniques that multiple paths flow into (e.g., both initial access vectors lead to the same execution technique)
- Identify BRANCH POINTS — techniques that lead to multiple downstream techniques
- If the report describes only one path, return a single path. Do not invent paths not supported by the report.
- Use only technique IDs that appear in refined_iocs or additional_mitre_techniques

CRITICAL — TACTIC ASSIGNMENT IN ATTACK PATHS:
Each step's tactic MUST reflect the technique's ROLE AT THAT POSITION in the path, not its most common tactic. "Initial Access" should only appear at the FIRST step of a path — it means the entry point. If the same technique appears later (after access was already gained), assign the tactic that matches its role at that later stage.

Many techniques have multiple valid MITRE tactics. Common multi-tactic techniques:
- T1078 (Valid Accounts): Initial Access, Persistence, Privilege Escalation, Defense Evasion
- T1053 (Scheduled Task/Job): Execution, Persistence, Privilege Escalation
- T1098 (Account Manipulation): Persistence, Privilege Escalation

Example — T1078 reused across two paths with DIFFERENT tactics:
  Path "cred-spray": T1110 (Credential Access) → T1078 (tactic: "Initial Access") → T1087 (Discovery)
    ↑ T1078 IS the initial foothold here — correct tactic is Initial Access.

  Path "aitm-phishing": T1566 (Initial Access) → T1557 (Credential Access) → T1078 (tactic: "Persistence")
    ↑ Initial access already happened at T1566. T1078 is using stolen creds to MAINTAIN access — correct tactic is Persistence.

  Path "insider-escalation": T1078 (tactic: "Initial Access") → T1068 (Privilege Escalation) → T1078 (tactic: "Privilege Escalation")
    ↑ Same technique at two different stages — different roles, different tactics.

SOURCE CONTEXT: {source_context}

IOC CANDIDATES:
{ioc_candidates}

FULL REPORT TEXT (truncated):
{report_text}

Respond ONLY with a JSON object in this exact format:
{{
  "refined_iocs": [
    {{
      "value": "the IOC value",
      "ioc_type": "ip|domain|hash_sha256|etc",
      "confidence": "low|medium|high",
      "priority": "low|medium|high",
      "context": "brief description of the IOC's role",
      "related_mitre": ["T1234", "T1234.001"],
      "is_false_positive": false
    }}
  ],
  "additional_mitre_techniques": [
    {{
      "technique_id": "T1234.001",
      "technique_name": "Technique Name",
      "tactic": "Tactic Name",
      "confidence": "explicit|inferred",
      "report_context": "brief description of the behavior"
    }},
    {{
      "technique_id": "T1078",
      "technique_name": "Valid Accounts",
      "tactic": "Initial Access",
      "confidence": "inferred",
      "report_context": "Attacker used harvested credentials for initial foothold"
    }},
    {{
      "technique_id": "T1078",
      "technique_name": "Valid Accounts",
      "tactic": "Persistence",
      "confidence": "inferred",
      "report_context": "Same credentials reused to maintain long-term access"
    }}
  ],
  "report_metadata": {{
    "title": "report title if identifiable",
    "campaign_name": "named campaign if mentioned",
    "attributed_actor": "threat actor if attributed",
    "attribution_confidence": "low|medium|high"
  }},
  "attack_graph": {{
    "paths": [
      {{
        "path_id": "short-slug-name",
        "description": "One sentence describing this attack path",
        "steps": [
          {{
            "technique_id": "T1566.001",
            "tactic": "Initial Access",
            "leads_to": ["T1059.001"]
          }},
          {{
            "technique_id": "T1059.001",
            "tactic": "Execution",
            "leads_to": ["T1053.005", "T1210"]
          }}
        ]
      }}
    ],
    "convergence_points": ["T1059.001"],
    "branch_points": ["T1059.001"]
  }}
}}
"""


# ToolResult, ValidationResult, QueryHints imported from framework.plugins.protocol


# ---------------------------------------------------------------------------
# Tool Implementation
# ---------------------------------------------------------------------------


class ThreatIntelIngester:
    """Event Mill plugin: Threat Intelligence Report Ingester.

    Extracts IOCs from threat intelligence reports (PDF, HTML, text)
    using a two-pass approach:
    1. Regex-based extraction for baseline IOC identification
    2. LLM-based refinement for confidence scoring, false positive
       filtering, MITRE ATT&CK mapping, and priority assessment
    """

    def __init__(self) -> None:
        self._manifest: dict | None = None

    def _load_manifest(self) -> dict:
        if self._manifest is None:
            manifest_path = Path(__file__).parent / "manifest.json"
            with open(manifest_path) as f:
                self._manifest = json.load(f)
        return self._manifest

    def metadata(self) -> dict:
        """Return runtime metadata reflecting the manifest."""
        manifest = self._load_manifest()
        return {
            "tool_name": manifest["tool_name"],
            "version": manifest["version"],
            "pillar": manifest["pillar"],
            "display_name": manifest["display_name"],
            "description_short": manifest["description_short"],
            "stability": manifest["stability"],
            "requires_llm": manifest["requires_llm"],
            "artifacts_consumed": manifest["artifacts_consumed"],
            "artifacts_produced": manifest["artifacts_produced"],
        }

    def validate_inputs(self, payload: dict) -> ValidationResult:
        """Validate the input payload against the input schema."""
        errors = []

        if "artifact_id" not in payload:
            errors.append("artifact_id is required")

        if "ioc_types" in payload:
            valid_types = {
                "ip", "domain", "hash_md5", "hash_sha1", "hash_sha256",
                "url", "email", "cve", "mitre_technique",
            }
            for t in payload["ioc_types"]:
                if t not in valid_types:
                    errors.append(f"Unknown ioc_type: {t}")

        if "confidence_threshold" in payload:
            if payload["confidence_threshold"] not in ("low", "medium", "high"):
                errors.append("confidence_threshold must be low, medium, or high")

        if "max_pages" in payload:
            mp = payload["max_pages"]
            if not isinstance(mp, int) or mp < 1 or mp > 200:
                errors.append("max_pages must be an integer between 1 and 200")

        return ValidationResult(ok=len(errors) == 0, errors=errors if errors else None)

    def execute(self, payload: dict, context: Any) -> ToolResult:
        """Ingest a threat intelligence report and extract structured IOC data.

        Two-pass extraction:
        1. Regex pass: identify IOC candidates from raw text
        2. LLM pass: refine confidence, filter false positives, map MITRE techniques
        """
        artifact_id = payload.get("artifact_id")
        if not artifact_id:
            return ToolResult(
                ok=False,
                error_code="INPUT_VALIDATION_FAILED",
                message=(
                    "artifact_id is required. "
                    "Usage: run threat_intel_ingester {\"artifact_id\": \"<id>\"}"
                ),
            )
        source_context = payload.get("source_context", "")
        ioc_types = payload.get(
            "ioc_types",
            ["ip", "domain", "hash_sha256", "url", "cve", "mitre_technique"],
        )
        confidence_threshold = payload.get("confidence_threshold", "low")
        max_pages = payload.get("max_pages", 50)

        # --- Resolve artifact ---
        artifact = None
        for art in context.artifacts:
            if art.artifact_id == artifact_id:
                artifact = art
                break

        if artifact is None:
            return ToolResult(
                ok=False,
                error_code="ARTIFACT_NOT_FOUND",
                message=(
                    f"Artifact {artifact_id!r} not found in session. "
                    f"Use 'artifacts' to list loaded artifacts."
                ),
            )

        if artifact.artifact_type not in ("pdf_report", "html_report", "text", "docx_report"):
            return ToolResult(
                ok=False,
                error_code="INPUT_VALIDATION_FAILED",
                message=(
                    f"Artifact type '{artifact.artifact_type}' is not supported. "
                    f"Expected pdf_report, html_report, text (including .md), or docx_report."
                ),
            )

        # --- Extract text ---
        logger.info(
            "Extracting text from %s artifact %s",
            artifact.artifact_type,
            artifact_id,
        )
        extractor = TEXT_EXTRACTORS[artifact.artifact_type]
        try:
            if artifact.artifact_type == "pdf_report":
                raw_text = extractor(artifact.file_path, max_pages)
            else:
                raw_text = extractor(str(artifact.file_path))
        except Exception as e:
            logger.error("Text extraction failed: %s", e)
            return ToolResult(
                ok=False,
                error_code="ARTIFACT_UNREADABLE",
                message=f"Failed to extract text: {e}",
            )

        if artifact.artifact_type == "pdf_report":
            page_count = raw_text.count("\n\n") + 1
        else:
            page_count = len(raw_text.splitlines())

        # --- Regex extraction pass ---
        logger.info("Running regex IOC extraction for types: %s", ioc_types)
        raw_iocs = extract_iocs_regex(raw_text, ioc_types)
        logger.info("Regex pass found %d IOC candidates", len(raw_iocs))

        # --- LLM refinement pass ---
        refined_iocs = []
        mitre_mappings = []
        report_meta = {}
        attack_graph = {}  # multi-path attack graph from LLM

        if context.llm_enabled and context.llm_query is not None:
            logger.info("[DIAG] LLM enabled, llm_query type=%s", type(context.llm_query).__name__)
            logger.info("Running LLM refinement on %d IOC candidates", len(raw_iocs))

            # Use framework reference data for MITRE grounding
            grounding: list[str] = []
            if hasattr(context, "reference_data"):
                mitre_data = context.reference_data.get("mitre_attack_enterprise")
                if mitre_data:
                    grounding.append(
                        "MITRE ATT&CK Enterprise techniques are available "
                        "for validation. Use official technique IDs."
                    )

            # --- Native PDF path: send full document to LLM directly ---
            native_pdf_succeeded = False
            if (artifact.artifact_type == "pdf_report"
                    and hasattr(context.llm_query, 'supports_native_document')
                    and context.llm_query.supports_native_document("application/pdf")):
                logger.info(
                    "Native PDF ingestion available — attempting "
                    "query_with_document()"
                )
                candidates_text = (
                    "\n".join(
                        f"- [{ioc.ioc_type}] {ioc.value} "
                        f"| Context: {ioc.context[:200]}"
                        for ioc in raw_iocs
                    )
                    or "(none found by regex pre-scan)"
                )
                native_prompt = LLM_REFINEMENT_PROMPT.format(
                    source_context=source_context or "Not provided",
                    ioc_candidates=candidates_text,
                    report_text=(
                        "[Full PDF document is attached — analyze the "
                        "complete document directly instead of this "
                        "placeholder text.]"
                    ),
                )
                try:
                    native_response = context.llm_query.query_with_document(
                        prompt=native_prompt,
                        artifact=artifact,
                        system_context=(
                            "You are a threat intelligence analyst. "
                            "Respond only with valid JSON."
                        ),
                        max_tokens=16384,
                        grounding_data=grounding,
                        hints=QueryHints(
                            tier="heavy",
                            prefers_native_file=True,
                            needs_structured_output=True,
                        ),
                    )
                    log_llm_interaction(
                        prompt=(
                            f"[ti_ingester native_pdf] "
                            f"{native_prompt[:500]}"
                        ),
                        response_text=native_response.text,
                        model_id=(
                            native_response.model_used
                            or "threat_intel_ingester"
                        ),
                        error=(
                            str(native_response.error)
                            if not native_response.ok else None
                        ),
                    )
                    if native_response.ok and native_response.text:
                        parsed = _parse_llm_json(native_response.text)
                        if parsed:
                            logger.info(
                                "Native PDF ingestion succeeded "
                                "(transport=%s, model=%s)",
                                native_response.transport_path,
                                native_response.model_used,
                            )
                            all_refined = parsed.get("refined_iocs", [])
                            refined_iocs = [
                                r for r in all_refined
                                if not r.get("is_false_positive", False)
                            ]
                            mitre_mappings = parsed.get(
                                "additional_mitre_techniques", [],
                            )
                            report_meta = parsed.get(
                                "report_metadata", {},
                            )
                            attack_graph = parsed.get(
                                "attack_graph", {},
                            )
                            native_pdf_succeeded = True
                        else:
                            logger.warning(
                                "Native PDF JSON parse failed — "
                                "falling back to chunked text path "
                                "| response_length=%d, model=%s, "
                                "transport=%s, token_usage=%s, "
                                "first_200=%r, last_200=%r",
                                len(native_response.text),
                                native_response.model_used,
                                native_response.transport_path,
                                native_response.token_usage,
                                native_response.text[:200],
                                native_response.text[-200:],
                            )
                            # Also log to activity so it shows in GCP
                            log_llm_interaction(
                                prompt="[ti_ingester native_pdf] JSON_PARSE_FAILED",
                                response_text=native_response.text,
                                model_id=(
                                    native_response.model_used
                                    or "threat_intel_ingester"
                                ),
                                error=(
                                    f"JSON parse failed on {len(native_response.text)}-char "
                                    f"response. first_100={native_response.text[:100]!r}"
                                ),
                            )
                    else:
                        logger.warning(
                            "Native PDF query returned failure — "
                            "falling back to chunked text path "
                            "| ok=%s, error=%r, model=%s, "
                            "transport=%s, fallback_reason=%s, "
                            "has_text=%s, text_length=%d",
                            native_response.ok,
                            native_response.error,
                            native_response.model_used,
                            native_response.transport_path,
                            native_response.fallback_reason,
                            native_response.text is not None,
                            len(native_response.text or ""),
                        )
                except Exception as e:
                    logger.error(
                        "Native PDF path exception — "
                        "falling back to chunked text path "
                        "| exception_type=%s, message=%s",
                        type(e).__name__, e,
                        exc_info=True,
                    )

            # Split large inputs into chunks to stay within LLM context limits.
            # Large PDFs previously caused repeated 503s even with backoff because
            # the monolithic prompt (100 IOCs + 8 kB text) exceeded the model's
            # comfortable input window.  Each chunk call is ~7-10 kB total.
            ioc_batches = [
                raw_iocs[i:i + _MAX_IOC_PER_CHUNK]
                for i in range(0, max(1, len(raw_iocs)), _MAX_IOC_PER_CHUNK)
            ]
            text_chunks = _chunk_text(raw_text, _MAX_TEXT_CHARS_PER_CHUNK)
            n_chunks = max(len(ioc_batches), len(text_chunks))

            if native_pdf_succeeded:
                n_chunks = 0
                logger.info(
                    "Skipping chunked LLM path — native PDF "
                    "ingestion succeeded"
                )

            logger.info(
                "Splitting into %d LLM chunk(s) (%d IOC batches, %d text chunks)",
                n_chunks, len(ioc_batches), len(text_chunks),
            )

            chunk_results: list[dict] = []
            chunk_json_failures = 0
            chunk_llm_failures = 0
            chunk_exceptions = 0
            for i in range(n_chunks):
                ioc_batch = ioc_batches[i] if i < len(ioc_batches) else []
                text_chunk = text_chunks[i] if i < len(text_chunks) else ""

                if not ioc_batch and not text_chunk:
                    continue

                candidates_text = (
                    "\n".join(
                        f"- [{ioc.ioc_type}] {ioc.value} | Context: {ioc.context[:200]}"
                        for ioc in ioc_batch
                    )
                    or "(none in this section)"
                )

                prompt = LLM_REFINEMENT_PROMPT.format(
                    source_context=source_context or "Not provided",
                    ioc_candidates=candidates_text,
                    report_text=text_chunk,
                )

                logger.info(
                    "LLM chunk %d/%d — %d IOCs, %d chars text",
                    i + 1, n_chunks, len(ioc_batch), len(text_chunk),
                )
                logger.info(
                    "[DIAG] Chunk %d/%d PROMPT SENT (%d chars). "
                    "First 500: %.500s",
                    i + 1, n_chunks, len(prompt), prompt[:500],
                )

                try:
                    llm_response = context.llm_query.query_text(
                        prompt=prompt,
                        system_context=(
                            "You are a threat intelligence analyst. "
                            "Respond only with valid JSON."
                        ),
                        max_tokens=4096,
                        grounding_data=grounding,
                        hints=QueryHints(
                            tier="light",
                            needs_structured_output=True,
                        ),
                    )

                    # --- Cloud Logging audit: input + output ---
                    log_llm_interaction(
                        prompt=f"[ti_ingester chunk {i+1}/{n_chunks}] {prompt[:500]}",
                        response_text=llm_response.text,
                        model_id=llm_response.model_used or "threat_intel_ingester",
                        error=(
                            str(llm_response.error)
                            if not llm_response.ok else None
                        ),
                    )

                    if llm_response.ok and llm_response.text:
                        logger.info(
                            "[DIAG] Chunk %d/%d LLM RESPONSE (%d chars). "
                            "First 500: %.500s",
                            i + 1, n_chunks, len(llm_response.text),
                            llm_response.text[:500],
                        )
                        parsed = _parse_llm_json(llm_response.text)
                        if parsed:
                            n_refined = len(parsed.get("refined_iocs", []))
                            n_fp = sum(
                                1 for r in parsed.get("refined_iocs", [])
                                if r.get("is_false_positive")
                            )
                            n_mitre = len(parsed.get("additional_mitre_techniques", []))
                            n_paths = len(parsed.get("attack_graph", {}).get("paths", []))
                            logger.info(
                                "[DIAG] Chunk %d/%d parsed OK — "
                                "%d refined_iocs (%d false_pos), "
                                "%d additional_mitre, %d attack_paths",
                                i + 1, n_chunks,
                                n_refined, n_fp, n_mitre, n_paths,
                            )
                            chunk_results.append(parsed)
                        else:
                            chunk_json_failures += 1
                            logger.warning(
                                "[DIAG] Chunk %d/%d: JSON parse FAILED "
                                "| response_length=%d, model=%s, "
                                "first_200=%r, last_200=%r",
                                i + 1, n_chunks,
                                len(llm_response.text),
                                llm_response.model_used,
                                llm_response.text[:200],
                                llm_response.text[-200:],
                            )
                    else:
                        logger.warning(
                            "[DIAG] Chunk %d/%d LLM call failed: ok=%s, "
                            "has_text=%s, error=%s",
                            i + 1, n_chunks, llm_response.ok,
                            bool(llm_response.text), llm_response.error,
                        )
                        chunk_llm_failures += 1
                        logger.warning(
                            "[DIAG] Chunk %d/%d FAILED RESPONSE "
                            "| model=%s, error=%r, "
                            "has_text=%s, text_length=%d, "
                            "first_200=%r",
                            i + 1, n_chunks,
                            llm_response.model_used,
                            llm_response.error,
                            llm_response.text is not None,
                            len(llm_response.text or ""),
                            (llm_response.text or "")[:200],
                        )

                except Exception as e:
                    chunk_exceptions += 1
                    logger.error(
                        "[DIAG] Chunk %d/%d EXCEPTION "
                        "| type=%s, message=%s",
                        i + 1, n_chunks, type(e).__name__, e,
                        exc_info=True,
                    )

            logger.info(
                "[DIAG] LLM loop done — %d/%d chunks produced parseable JSON",
                len(chunk_results), n_chunks,
            )

            if chunk_results:
                merged = _merge_llm_chunk_results(chunk_results)
                all_refined = merged.get("refined_iocs", [])
                non_fp = [r for r in all_refined if not r.get("is_false_positive", False)]
                logger.info(
                    "[DIAG] Merged result — %d refined_iocs total, "
                    "%d after false-positive filter, "
                    "%d additional_mitre, attack_graph paths=%d",
                    len(all_refined), len(non_fp),
                    len(merged.get("additional_mitre_techniques", [])),
                    len(merged.get("attack_graph", {}).get("paths", [])),
                )
                for refined in all_refined:
                    if not refined.get("is_false_positive", False):
                        refined_iocs.append(refined)
                mitre_mappings = merged.get("additional_mitre_techniques", [])
                report_meta = merged.get("report_metadata", {})
                attack_graph = merged.get("attack_graph", {})
            elif not native_pdf_succeeded:
                logger.warning(
                    "[DIAG] All %d LLM chunks failed — "
                    "json_parse_failures=%d, llm_call_failures=%d, "
                    "exceptions=%d — falling back to regex-only.",
                    n_chunks, chunk_json_failures,
                    chunk_llm_failures, chunk_exceptions,
                )

        # If LLM refinement didn't produce results, use regex baseline
        ingestion_mode = "llm"  # track which path produced results
        if not refined_iocs:
            ingestion_mode = "regex_only"
            llm_was_enabled = context.llm_enabled and context.llm_query is not None
            logger.warning(
                "[DIAG] FALLBACK to regex-only — refined_iocs empty. "
                "LLM enabled=%s, mitre_mappings=%d, attack_graph_paths=%d",
                llm_was_enabled, len(mitre_mappings),
                len(attack_graph.get("paths", [])) if isinstance(attack_graph, dict) else 0,
            )
            logger.info("Using regex-only IOC results (no LLM refinement)")
            refined_iocs = [
                {
                    "ioc_type": ioc.ioc_type,
                    "value": ioc.value,
                    "confidence": "low",
                    "priority": "medium",
                    "context": ioc.context[:300],
                    "related_mitre": [],
                    "defanged": ioc.defanged,
                }
                for ioc in raw_iocs
            ]

        # --- Apply confidence threshold filter ---
        confidence_order = {"low": 0, "medium": 1, "high": 2}
        threshold_value = confidence_order.get(confidence_threshold, 0)
        filtered_iocs = [
            ioc
            for ioc in refined_iocs
            if confidence_order.get(ioc.get("confidence", "low"), 0) >= threshold_value
        ]

        # --- Build MITRE mappings from IOCs + additional techniques ---
        all_mitre = list(mitre_mappings)  # Start with additional techniques
        seen_techniques = {m["technique_id"] for m in all_mitre}

        for ioc in filtered_iocs:
            for tech_id in ioc.get("related_mitre", []):
                if tech_id not in seen_techniques:
                    seen_techniques.add(tech_id)
                    all_mitre.append(
                        {
                            "technique_id": tech_id,
                            "technique_name": "",  # enriched by _reconcile below
                            "tactic": "",
                            "confidence": "inferred",
                            "report_context": f"Associated with IOC {ioc['value']}",
                        }
                    )

        # --- Reconcile: backfill + enrich from local ATT&CK data ---
        all_mitre = _reconcile_mitre_mappings(all_mitre, attack_graph)

        # --- Build summary ---
        ioc_breakdown: dict[str, int] = {}
        high_priority_count = 0
        confidence_dist = {"low": 0, "medium": 0, "high": 0}

        for ioc in filtered_iocs:
            ioc_type = ioc["ioc_type"]
            ioc_breakdown[ioc_type] = ioc_breakdown.get(ioc_type, 0) + 1
            if ioc.get("priority") == "high":
                high_priority_count += 1
            conf = ioc.get("confidence", "low")
            confidence_dist[conf] = confidence_dist.get(conf, 0) + 1

        # --- Register output artifact ---
        output_artifact_path = None
        output_artifact_id = None

        if hasattr(context, "register_artifact") and context.register_artifact:
            output_data = {
                "report_metadata": report_meta,
                "iocs": filtered_iocs,
                "mitre_mappings": all_mitre,
                "attack_graph": attack_graph,
            }

            # Write artifact file
            import tempfile
            import os

            workspace = os.environ.get("EVENTMILL_WORKSPACE", "/tmp")
            artifact_dir = os.path.join(workspace, "artifacts")
            os.makedirs(artifact_dir, exist_ok=True)

            with tempfile.NamedTemporaryFile(
                mode="w",
                suffix="_ti_iocs.json",
                dir=artifact_dir,
                delete=False,
                prefix="art_",
            ) as f:
                json.dump(output_data, f, indent=2)
                output_artifact_path = f.name

            art_ref = context.register_artifact(
                artifact_type="json_events",
                file_path=output_artifact_path,
                source_tool="threat_intel_ingester",
                metadata={
                    "ioc_count": len(filtered_iocs),
                    "source_artifact": artifact_id,
                },
            )
            output_artifact_id = art_ref.artifact_id

        # --- Build result ---
        output_artifacts_list = None
        if output_artifact_id:
            output_artifacts_list = [
                {
                    "artifact_id": output_artifact_id,
                    "artifact_type": "json_events",
                    "file_path": output_artifact_path,
                    "description": "Structured IOC records extracted from threat intel report.",
                }
            ]

        logger.info(
            "Ingestion complete: %d IOCs, %d MITRE techniques, %d high-priority",
            len(filtered_iocs),
            len(all_mitre),
            high_priority_count,
        )

        return ToolResult(
            ok=True,
            result={
                "report_metadata": {
                    "title": report_meta.get("title", ""),
                    "source_organization": report_meta.get("source_organization", ""),
                    "publication_date": report_meta.get("publication_date", ""),
                    "page_count": page_count,
                    "artifact_type": artifact.artifact_type,
                    "campaign_name": report_meta.get("campaign_name", ""),
                    "attributed_actor": report_meta.get("attributed_actor", ""),
                    "attribution_confidence": report_meta.get(
                        "attribution_confidence", ""
                    ),
                },
                "iocs": filtered_iocs,
                "mitre_mappings": all_mitre,
                "attack_graph": attack_graph,
                "summary": {
                    "total_iocs": len(filtered_iocs),
                    "ioc_breakdown": ioc_breakdown,
                    "high_priority_count": high_priority_count,
                    "mitre_technique_count": len(all_mitre),
                    "unique_technique_count": len(
                        {m.get("technique_id") for m in all_mitre
                         if m.get("technique_id")}
                    ),
                    "confidence_distribution": confidence_dist,
                    "ingestion_mode": ingestion_mode,
                },
            },
            output_artifacts=output_artifacts_list,
        )

    def summarize_for_llm(self, result: Any) -> str:
        """Produce a compressed summary for LLM context window."""
        if not result.ok:
            error = result.message or "Unknown error"
            return f"Threat intel ingestion failed: {error}"

        r = result.result or {}
        meta = r.get("report_metadata", {})
        summary = r.get("summary", {})
        mitre = r.get("mitre_mappings", [])

        parts = []

        # Report identity
        title = meta.get("title", "Unknown report")
        artifact_type = meta.get("artifact_type", "unknown")
        pages = meta.get("page_count", "?")
        size_label = "pages" if artifact_type == "pdf_report" else "lines"
        parts.append(f"Ingested {artifact_type} ({pages} {size_label}): {title}.")

        # Attribution
        actor = meta.get("attributed_actor")
        campaign = meta.get("campaign_name")
        if actor:
            conf = meta.get("attribution_confidence", "")
            parts.append(
                f"Attributed to {actor}"
                + (f" ({conf} confidence)" if conf else "")
                + (f", campaign: {campaign}" if campaign else "")
                + "."
            )

        # IOC counts
        total = summary.get("total_iocs", 0)
        breakdown = summary.get("ioc_breakdown", {})
        breakdown_str = ", ".join(
            f"{count} {ioc_type}{'s' if count != 1 else ''}"
            for ioc_type, count in sorted(breakdown.items())
        )
        parts.append(f"Extracted {total} IOCs: {breakdown_str}.")

        # High priority
        hp = summary.get("high_priority_count", 0)
        if hp > 0:
            parts.append(f"{hp} IOCs flagged as high-priority.")

        # MITRE techniques — group by technique_id for multi-role display
        tech_count = summary.get("mitre_technique_count", 0)
        unique_count = summary.get("unique_technique_count", tech_count)
        if tech_count > 0 and mitre:
            by_tid: dict[str, list[str]] = {}
            names: dict[str, str] = {}
            for m in mitre:
                tid = m.get("technique_id", "")
                if not tid:
                    continue
                by_tid.setdefault(tid, [])
                tac = m.get("tactic", "")
                if tac and tac not in by_tid[tid]:
                    by_tid[tid].append(tac)
                if not names.get(tid):
                    names[tid] = m.get("technique_name", "")

            tech_parts = []
            for tid in list(by_tid.keys())[:6]:
                tactics = by_tid[tid]
                if len(tactics) > 1:
                    tech_parts.append(f"{tid} ({', '.join(tactics)})")
                else:
                    tech_parts.append(f"{tid} ({names.get(tid, '')})")

            if unique_count != tech_count:
                parts.append(
                    f"Mapped to {unique_count} unique techniques across "
                    f"{tech_count} tactical roles: {', '.join(tech_parts)}."
                )
            else:
                parts.append(
                    f"Mapped to {tech_count} MITRE techniques: "
                    f"{', '.join(tech_parts)}."
                )
            if len(by_tid) > 6:
                parts.append(f"(and {len(by_tid) - 6} more)")

        # Attack graph paths
        if r.get("attack_graph", {}).get("paths"):
            path_count = len(r["attack_graph"]["paths"])
            convergence = r["attack_graph"].get("convergence_points", [])
            parts.append(
                f"Attack graph: {path_count} path(s) identified"
                + (f", converging at {', '.join(convergence)}" if convergence else "")
                + "."
            )

        # Ingestion mode warning
        mode = summary.get("ingestion_mode", "")
        if mode == "regex_only":
            parts.append(
                "WARNING: LLM analysis failed — results are regex-only "
                "(low confidence, no MITRE mapping, no attack graph). "
                "Check logs for LLM failure details."
            )

        # Output artifact + quick chart command
        artifacts = result.output_artifacts or []
        if artifacts:
            art = artifacts[0]
            aid = art['artifact_id']
            parts.append(
                f"Output artifact: {aid} "
                f"({art['artifact_type']})."
            )
            parts.append(
                f"Quick chart: run attack_path_visualizer "
                f'{{\"artifact_id\": \"{aid}\", \"format\": \"mermaid\"}}'
            )

        return " ".join(parts)
